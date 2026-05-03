import sys
import os
import shutil
import subprocess
import tempfile
from datetime import datetime

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QFileDialog, QLabel, QApplication, QFrame,
    QFileIconProvider, QScrollArea, QSizePolicy,
    QDialog, QComboBox
)
from PyQt5.QtCore import Qt, QFileInfo, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QPixmap

IMAGE_EXTENSIONS  = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".ico"}
VIDEO_EXTENSIONS  = {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm"}
AUDIO_EXTENSIONS  = {".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a", ".opus"}
MISC_EXTENSIONS   = {".pdf", ".docx", ".xlsx", ".pptx", ".txt", ".zip", ".rar", ".7z"}

CONVERSION_TARGETS = {
    "image": [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"],
    "video": [".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".gif"],
    "audio": [".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a", ".opus"],
    "misc":  [".pdf", ".txt"],
}

def get_category(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext in IMAGE_EXTENSIONS: return "image"
    if ext in VIDEO_EXTENSIONS:  return "video"
    if ext in AUDIO_EXTENSIONS:  return "audio"
    return "misc"

def find_ffmpeg():
    if getattr(sys, "frozen", False):
        bundled = os.path.join(sys._MEIPASS, "ffmpeg", "bin", "ffmpeg.exe")
        if os.path.isfile(bundled):
            return bundled

    p = shutil.which("ffmpeg")
    if p:
        return p

    candidates = [
        r"C:\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files (x86)\ffmpeg\bin\ffmpeg.exe",
        os.path.join(os.path.expanduser("~"), "ffmpeg", "bin", "ffmpeg.exe"),
    ]
    for c in candidates:
        if os.path.isfile(c):
            return c

    import glob
    patterns = [
        r"C:\ffmpeg*\bin\ffmpeg.exe",
        r"C:\ffmpeg*\ffmpeg*\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg*\bin\ffmpeg.exe",
        r"C:\Users\*\ffmpeg*\bin\ffmpeg.exe",
        r"C:\Users\*\Downloads\ffmpeg*\bin\ffmpeg.exe",
    ]
    for pattern in patterns:
        matches = glob.glob(pattern)
        if matches:
            return matches[0]

    return None

FFMPEG = find_ffmpeg()

def safe_path(p):
    return os.path.normpath(os.path.abspath(p))


class ConvertWorker(QThread):
    finished = pyqtSignal(str)
    error    = pyqtSignal(str)

    def __init__(self, src, target_ext):
        super().__init__()
        self.src = src
        self.target_ext = target_ext

    def run(self):
        src = safe_path(self.src)
        ext = self.target_ext
        category = get_category(src)
        base = os.path.splitext(os.path.basename(src))[0]
        out_dir = tempfile.mkdtemp(prefix="fshift_")
        out_path = safe_path(os.path.join(out_dir, base + ext))

        try:
            if category == "image":
                self._convert_image(src, out_path, ext)
            elif category in ("video", "audio"):
                self._convert_ffmpeg(src, out_path)
            else:
                self._convert_misc(src, out_path, ext)
            self.finished.emit(out_path)
        except Exception as e:
            self.error.emit(str(e))

    def _run_ffmpeg(self, args):
        if not FFMPEG:
            raise RuntimeError(
                "ffmpeg not found. Install ffmpeg and ensure it is on your PATH.\n"
                "Download from https://ffmpeg.org/download.html"
            )
        flags = {}
        if sys.platform == "win32":
            flags["creationflags"] = subprocess.CREATE_NO_WINDOW
        result = subprocess.run(
            [FFMPEG] + args,
            capture_output=True,
            text=True,
            **flags
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr[-800:])

    def _convert_image(self, src, out_path, ext):
        from PIL import Image
        img = Image.open(src)
        if ext in (".jpg", ".jpeg") and img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.save(out_path)

    def _convert_ffmpeg(self, src, out_path):
        if out_path.lower().endswith(".gif"):
            self._convert_to_gif(src, out_path)
        else:
            self._run_ffmpeg(["-y", "-i", src, out_path])

    def _convert_to_gif(self, src, out_path):
        out_dir = os.path.dirname(out_path)
        palette_path = safe_path(os.path.join(out_dir, "_palette_tmp.png"))
        try:
            self._run_ffmpeg([
                "-y", "-i", src,
                "-vf", "fps=15,scale=480:-2:flags=lanczos,palettegen=stats_mode=diff",
                palette_path,
            ])
            self._run_ffmpeg([
                "-y", "-i", src, "-i", palette_path,
                "-filter_complex",
                "fps=15,scale=480:-2:flags=lanczos[x];[x][1:v]paletteuse=dither=bayer",
                out_path,
            ])
        finally:
            if os.path.exists(palette_path):
                try:
                    os.remove(palette_path)
                except OSError:
                    pass

    def _convert_misc(self, src, out_path, ext):
        src_ext = os.path.splitext(src)[1].lower()
        if ext == ".pdf":
            self._to_pdf(src, out_path, src_ext)
        elif ext == ".txt":
            self._to_txt(src, out_path, src_ext)
        else:
            shutil.copy2(src, out_path)

    def _to_pdf(self, src, out_path, src_ext):
        if src_ext == ".txt":
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import letter
            c = rl_canvas.Canvas(out_path, pagesize=letter)
            w, h = letter
            y = h - 40
            with open(src, "r", errors="replace") as f:
                for line in f:
                    c.drawString(40, y, line.rstrip())
                    y -= 14
                    if y < 40:
                        c.showPage()
                        y = h - 40
            c.save()
        elif src_ext == ".docx":
            import docx
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import letter
            doc = docx.Document(src)
            c = rl_canvas.Canvas(out_path, pagesize=letter)
            w, h = letter
            y = h - 40
            for para in doc.paragraphs:
                c.drawString(40, y, para.text)
                y -= 14
                if y < 40:
                    c.showPage()
                    y = h - 40
            c.save()
        else:
            shutil.copy2(src, out_path)

    def _to_txt(self, src, out_path, src_ext):
        if src_ext == ".docx":
            import docx
            doc = docx.Document(src)
            with open(out_path, "w") as f:
                for para in doc.paragraphs:
                    f.write(para.text + "\n")
        elif src_ext == ".pdf":
            flags = {}
            if sys.platform == "win32":
                flags["creationflags"] = subprocess.CREATE_NO_WINDOW
            result = subprocess.run(
                ["pdftotext", src, out_path],
                capture_output=True, **flags
            )
            if result.returncode != 0:
                raise RuntimeError("pdftotext failed — install poppler-utils")
        else:
            shutil.copy2(src, out_path)


class FileTypeDialog(QDialog):
    def __init__(self, category, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Format")
        self.setFixedSize(250, 100)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowContextHelpButtonHint)

        layout = QVBoxLayout()
        self.label = QLabel("Choose output format:")
        layout.addWidget(self.label)
        self.combo = QComboBox()
        targets = CONVERSION_TARGETS.get(category, [".txt"])
        self.combo.addItems(targets)
        layout.addWidget(self.combo)
        self.btn = QPushButton("Convert")
        self.btn.clicked.connect(self.accept)
        layout.addWidget(self.btn)
        self.setLayout(layout)

    def get_selection(self):
        return self.combo.currentText()


STYLE = {
    "default": """
        * { font-family: 'JetBrains Mono'; outline: none; font-size: 12px; }
    """
}


class FileShiftUi(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("FileShift")
        self._loaded_files = []
        self._cards = []
        self.init_ui()

    def init_ui(self):
        self.resize(500, 800)

        self.master_layout = QVBoxLayout(self)
        self.master_layout.setContentsMargins(20, 10, 20, 10)
        self.master_layout.setSpacing(8)

        self.main_container = QWidget()
        self.main_container.setStyleSheet(STYLE["default"])
        self.master_layout.addWidget(self.main_container)

        root_layout = QVBoxLayout(self.main_container)
        root_layout.setContentsMargins(15, 15, 15, 15)
        root_layout.setSpacing(10)

        root_layout.addWidget(QLabel("Select files to convert:"))

        self.button_file = QPushButton("Select Files")
        self.button_file.clicked.connect(self.select_file)
        root_layout.addWidget(self.button_file)

        if not FFMPEG:
            warn = QLabel("! ffmpeg not found; some video/audio conversion will be unavailable.")
            warn.setStyleSheet("color: orange;")
            warn.setWordWrap(True)
            root_layout.addWidget(warn)

        self.files_container = QWidget()
        self.files_group_layout = QVBoxLayout(self.files_container)
        self.files_group_layout.setContentsMargins(0, 0, 0, 0)
        self.files_group_layout.setSpacing(8)
        self.files_group_layout.addStretch(1)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidget(self.files_container)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll_area.setFrameShape(QFrame.NoFrame)
        self.scroll_area.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        root_layout.addWidget(self.scroll_area, stretch=1)

        self.convert_all = QPushButton("Convert All")
        self.convert_all.setEnabled(False)
        self.convert_all.clicked.connect(self.convert_all_files)
        root_layout.addWidget(self.convert_all)

        self.download_all = QPushButton("Download All")
        self.download_all.setEnabled(False)
        self.download_all.clicked.connect(self.download_all_files)
        root_layout.addWidget(self.download_all)

        self.copywrite_display = QLabel("© 2026 Daniil Ovechkin. Built using Python and PyQt5.")
        self.copywrite_display.setAlignment(Qt.AlignCenter)
        root_layout.addWidget(self.copywrite_display)

    def get_icon_pixmap(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()
        if ext in IMAGE_EXTENSIONS:
            pixmap = QPixmap(filepath)
            if not pixmap.isNull():
                return pixmap.scaled(32, 32, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        provider = QFileIconProvider()
        icon = provider.icon(QFileInfo(filepath))
        if icon.isNull():
            icon = provider.icon(QFileIconProvider.File)
        return icon.pixmap(32, 32)

    def build_file_card(self, filepath):
        card = QWidget()
        card.setObjectName("card")
        card.setStyleSheet("QWidget#card { border: 1px solid #aaa; border-radius: 4px; }")
        card.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        card_layout = QHBoxLayout(card)
        card_layout.setContentsMargins(8, 8, 8, 8)
        card_layout.setSpacing(10)

        icon_label = QLabel()
        icon_label.setPixmap(self.get_icon_pixmap(filepath))
        icon_label.setFixedSize(32, 32)
        icon_label.setScaledContents(True)
        card_layout.addWidget(icon_label)

        text_widget = QWidget()
        text_layout = QVBoxLayout(text_widget)
        text_layout.setContentsMargins(0, 0, 0, 0)
        text_layout.setSpacing(2)

        file_name = os.path.basename(filepath)
        ext_upper = os.path.splitext(file_name)[1].lstrip(".").upper() or "FILE"
        size_bytes = os.path.getsize(filepath)
        if size_bytes < 1024:
            size_str = f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            size_str = f"{size_bytes / 1024:.1f} KB"
        else:
            size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
        mtime = datetime.fromtimestamp(os.path.getmtime(filepath)).strftime("%Y-%m-%d %H:%M")

        name_label = QLabel(file_name)
        name_label.setWordWrap(True)
        name_label.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Preferred)
        info_label = QLabel(f"{ext_upper} | {size_str} | {mtime}")

        text_layout.addWidget(name_label)
        text_layout.addWidget(info_label)
        card_layout.addWidget(text_widget, 1)

        convert_btn = QPushButton("Convert")
        convert_btn.clicked.connect(lambda: self.start_convert(filepath, card_info))
        card_layout.addWidget(convert_btn)

        status_label = QLabel("")
        status_label.setAlignment(Qt.AlignCenter)
        status_label.setVisible(False)
        card_layout.addWidget(status_label)

        dl_btn = QPushButton("Download")
        dl_btn.setVisible(False)
        card_layout.addWidget(dl_btn)

        close_btn = QPushButton("X")
        close_btn.setFixedSize(24, 24)
        close_btn.setStyleSheet("QPushButton { border: none; color: #888; } QPushButton:hover { color: #f00; }")
        close_btn.clicked.connect(lambda: self.remove_card(card_info))
        card_layout.addWidget(close_btn)

        card_info = {
            "filepath": filepath,
            "card": card,
            "convert_btn": convert_btn,
            "status_label": status_label,
            "dl_btn": dl_btn,
        }

        return card, card_info

    def _run_conversion(self, card_info, target_ext):
        filepath     = card_info["filepath"]
        card         = card_info["card"]
        convert_btn  = card_info["convert_btn"]
        status_label = card_info["status_label"]
        dl_btn       = card_info["dl_btn"]

        convert_btn.setVisible(False)
        dl_btn.setVisible(False)
        status_label.setText("Converting.")
        status_label.setVisible(True)

        dots = [0]
        timer = QTimer(self)
        def tick():
            dots[0] = (dots[0] + 1) % 4
            status_label.setText("Converting" + "." * dots[0])
        timer.timeout.connect(tick)
        timer.start(400)

        worker = ConvertWorker(filepath, target_ext)
        card._worker = worker
        card._timer  = timer

        def on_done(out_path):
            timer.stop()
            status_label.setVisible(False)
            dl_btn.setVisible(True)
            try:
                dl_btn.clicked.disconnect()
            except Exception:
                pass
            dl_btn.clicked.connect(lambda: self.download_file(out_path))
            card_info["out_path"] = out_path

        def on_error(msg):
            timer.stop()
            status_label.setText("Error!")
            status_label.setToolTip(msg)
            status_label.setVisible(True)

        worker.finished.connect(on_done)
        worker.error.connect(on_error)
        worker.start()

    def start_convert(self, filepath, card_info):
        category = get_category(filepath)
        dialog = FileTypeDialog(category, self)
        if not dialog.exec_():
            return
        target_ext = dialog.get_selection()
        self._run_conversion(card_info, target_ext)

    def convert_all_files(self):
        if not self._cards:
            return
        first_path = self._cards[0]["filepath"]
        category = get_category(first_path)
        dialog = FileTypeDialog(category, self)
        if not dialog.exec_():
            return
        target_ext = dialog.get_selection()
        for card_info in self._cards:
            if "out_path" not in card_info:
                self._run_conversion(card_info, target_ext)

    def download_all_files(self):
        ready = [ci for ci in self._cards if "out_path" in ci]
        if not ready:
            return
        dest_dir = QFileDialog.getExistingDirectory(self, "Select Download Folder")
        if not dest_dir:
            return
        for card_info in ready:
            out_path = card_info["out_path"]
            filename = os.path.basename(out_path)
            if not os.path.splitext(filename)[1]:
                filename += os.path.splitext(out_path)[1]
            dest = os.path.join(dest_dir, filename)
            shutil.copy2(out_path, dest)

    def download_file(self, out_path):
        ext = os.path.splitext(out_path)[1]
        dest, _ = QFileDialog.getSaveFileName(
            self,
            "Save File",
            os.path.basename(out_path),
            f"*{ext};;All Files (*)"
        )
        if dest:
            if not dest.lower().endswith(ext.lower()):
                dest += ext
            shutil.copy2(out_path, dest)

    def select_file(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Select Files")
        if not files:
            return
        files = [f for f in files if os.path.isfile(f)]
        added = 0
        for filepath in files:
            abs_path = os.path.abspath(filepath)
            if abs_path in self._loaded_files:
                continue
            self._loaded_files.append(abs_path)
            card, card_info = self.build_file_card(abs_path)
            self._cards.append(card_info)
            self.files_group_layout.insertWidget(self.files_group_layout.count() - 1, card)
            added += 1
        if added:
            self.convert_all.setEnabled(True)
            self.download_all.setEnabled(True)
        print(f"Selected files: {files} ({added} new)")

    def remove_card(self, card_info):
        abs_path = os.path.abspath(card_info["filepath"])
        if abs_path in self._loaded_files:
            self._loaded_files.remove(abs_path)
        if card_info in self._cards:
            self._cards.remove(card_info)
        card_info["card"].setParent(None)
        card_info["card"].deleteLater()
        if not self._cards:
            self.convert_all.setEnabled(False)
            self.download_all.setEnabled(False)